import os
import re
import sys
import shutil
import tempfile
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

try:
    from tkinterdnd2 import TkinterDnD, DND_FILES
    HAS_DND = True
except ImportError:
    HAS_DND = False

from PIL import Image, ImageTk, ImageOps
from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Constants ─────────────────────────────────────────────────────────────────
COL_W_EMU = 3_371_850
PAGE_W_DXA = 12240
PAGE_H_DXA = 15840
MARGIN_DXA = 720
TABLE_W_DXA = 10800
COL_W_DXA = 5310
CAPTION_SPACE_BEFORE_DXA = 40
CAPTION_SPACE_AFTER_DXA = 80
CAPTION_FONT_SIZE = 8
CAPTION_FONT_NAME = "Arial"
MAX_EMBED_PX = 1200  # compress images larger than this before embedding


# ── Utilities ─────────────────────────────────────────────────────────────────
def _resource_path(filename):
    base = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, filename)


def natural_sort_key(s):
    parts = re.split(r'(\d+)', s)
    return [int(p) if p.isdigit() else p.lower() for p in parts]


# ── DOCX helpers ──────────────────────────────────────────────────────────────
def set_cell_borders_none(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_margins_zero(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ('top', 'left', 'bottom', 'right'):
        mar = OxmlElement(f'w:{side}')
        mar.set(qn('w:w'), '0')
        mar.set(qn('w:type'), 'dxa')
        tcMar.append(mar)
    tcPr.append(tcMar)


def set_cell_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def set_table_borders_none(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'FFFFFF')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_table_width(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(TABLE_W_DXA))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def set_column_widths(table):
    tbl = table._tbl
    tblGrid = OxmlElement('w:tblGrid')
    for w in (COL_W_DXA, COL_W_DXA):
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    tbl.insert(1, tblGrid)


def add_caption(cell, text):
    # Reuse the cell's default first paragraph — avoids a spurious empty paragraph gap
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(CAPTION_SPACE_BEFORE_DXA))
    spacing.set(qn('w:after'), str(CAPTION_SPACE_AFTER_DXA))
    pPr.append(spacing)
    run = para.add_run(text)
    run.font.name = CAPTION_FONT_NAME
    run.font.size = Pt(CAPTION_FONT_SIZE)


def add_image_to_cell(cell, image_path, width_emu, height_emu):
    para = cell.paragraphs[0]
    pPr = para._p.get_or_add_pPr()
    # Suppress default Normal-style spacing-after so caption sits flush below photo
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    run = para.add_run()
    run.add_picture(image_path,
                    width=Inches(width_emu / 914400),
                    height=Inches(height_emu / 914400))


def get_image_dims(path):
    with Image.open(path) as img:
        img = ImageOps.exif_transpose(img)
        return img.width, img.height


def prepare_image(src_path, tmpdir, idx):
    """Normalize EXIF orientation and compress if needed. Always saves a copy."""
    with Image.open(src_path) as img:
        img = ImageOps.exif_transpose(img)  # apply rotation before anything else
        w, h = img.width, img.height
        if max(w, h) > MAX_EMBED_PX:
            ratio = MAX_EMBED_PX / max(w, h)
            img = img.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
        if img.mode in ('RGBA', 'P', 'LA'):
            img = img.convert('RGB')
        out_path = os.path.join(tmpdir, f'{idx:05d}.jpg')
        img.save(out_path, 'JPEG', quality=90)
    return out_path


def compute_row_sizes(row_images):
    """row_images: [(path, nat_w, nat_h), ...] → [(final_w_emu, final_h_emu), ...]"""
    if len(row_images) == 1:
        _, nat_w, nat_h = row_images[0]
        return [(COL_W_EMU, nat_h / nat_w * COL_W_EMU)]
    dhs = [nat_h / nat_w * COL_W_EMU for _, nat_w, nat_h in row_images]
    target_h = min(dhs)
    return [(target_h * nat_w / nat_h, target_h) for _, nat_w, nat_h in row_images]


# ── Core document builder ─────────────────────────────────────────────────────
def build_contact_sheet(image_paths, output_path, per_page, title,
                        progress_cb, status_cb):
    images = sorted(image_paths, key=lambda p: natural_sort_key(os.path.basename(p)))
    total = len(images)
    if total == 0:
        raise ValueError("No images selected.")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Twips(PAGE_W_DXA)
    section.page_height = Twips(PAGE_H_DXA)
    section.top_margin = Twips(MARGIN_DXA)
    section.bottom_margin = Twips(MARGIN_DXA)
    section.left_margin = Twips(MARGIN_DXA)
    section.right_margin = Twips(MARGIN_DXA)

    if title and title.strip():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = para._p.get_or_add_pPr()
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '120')
        pPr.append(sp)
        run = para.add_run(title.strip())
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.bold = True

    rows_per_page = per_page // 2
    pages = [images[i:i + per_page] for i in range(0, total, per_page)]
    first_page = True
    processed = 0
    embed_idx = 0
    tmpdir = tempfile.mkdtemp()

    try:
        for page_images in pages:
            if not first_page:
                pb_para = doc.add_paragraph()
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                pb_para.add_run()._r.append(br)
            first_page = False

            table = doc.add_table(rows=rows_per_page * 2, cols=2)
            set_table_borders_none(table)
            set_table_width(table)
            set_column_widths(table)

            for row_idx, row_imgs in enumerate(
                    [page_images[i:i + 2] for i in range(0, len(page_images), 2)]):
                row_info = [(p, *get_image_dims(p)) for p in row_imgs]
                sizes = compute_row_sizes(row_info)
                img_row = table.rows[row_idx * 2]
                cap_row = table.rows[row_idx * 2 + 1]

                for col_idx in range(2):
                    ic = img_row.cells[col_idx]
                    cc = cap_row.cells[col_idx]
                    for cell in (ic, cc):
                        set_cell_borders_none(cell)
                        set_cell_margins_zero(cell)
                        set_cell_width(cell, COL_W_DXA)
                    if col_idx < len(row_imgs):
                        path = row_imgs[col_idx]
                        fw, fh = sizes[col_idx]
                        embed = prepare_image(path, tmpdir, embed_idx)
                        embed_idx += 1
                        add_image_to_cell(ic, embed, fw, fh)
                        add_caption(cc, os.path.splitext(os.path.basename(path))[0])

                processed += len(row_imgs)
                progress_cb(processed / total * 100)
                status_cb(f"Processing image {processed} of {total}…")

        doc.save(output_path)
        status_cb(f"Done! Saved to {output_path}")
        progress_cb(100)
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


# ── GUI ───────────────────────────────────────────────────────────────────────
_AppBase = TkinterDnD.Tk if HAS_DND else tk.Tk


class App(_AppBase):
    def __init__(self):
        super().__init__()
        self.title("Contact Sheet Builder")
        self.resizable(False, False)
        self._image_order = []  # ordered list of absolute paths
        self._image_set = set() # fast duplicate check
        self._logo_img = None
        self._build_ui()
        if HAS_DND:
            self.drop_target_register(DND_FILES)
            self.dnd_bind('<<Drop>>', self._handle_drop)

    # ── UI construction ───────────────────────────────────────────────────────

    def _build_ui(self):
        pad = dict(padx=10, pady=5)

        # Logo
        logo_path = _resource_path('evoq_logo.png')
        if os.path.exists(logo_path):
            with Image.open(logo_path) as raw:
                dw = 280
                dh = round(raw.height / raw.width * dw)
                self._logo_img = ImageTk.PhotoImage(
                    raw.resize((dw, dh), Image.LANCZOS))
            tk.Label(self, image=self._logo_img, bg=self.cget('bg')).grid(
                row=0, column=0, columnspan=2, pady=(12, 4))

        # Images panel
        self._build_image_panel().grid(
            row=1, column=0, columnspan=2, sticky="ew", **pad)

        # Document settings
        settings = tk.LabelFrame(self, text="Document Settings")
        settings.grid(row=2, column=0, columnspan=2, sticky="ew", **pad)
        tk.Label(settings, text="Title (optional):").grid(
            row=0, column=0, sticky="w", padx=5, pady=4)
        self.title_var = tk.StringVar()
        tk.Entry(settings, textvariable=self.title_var, width=38).grid(
            row=0, column=1, sticky="ew", padx=5, pady=4)
        settings.columnconfigure(1, weight=1)

        # Output
        out_frame = tk.LabelFrame(self, text="Output File")
        out_frame.grid(row=3, column=0, columnspan=2, sticky="ew", **pad)
        docs = os.path.join(os.path.expanduser("~"), "Documents")
        default_dir = docs if os.path.isdir(docs) else os.path.expanduser("~")
        self.output_var = tk.StringVar(value=os.path.join(default_dir, "ContactSheet.docx"))
        tk.Entry(out_frame, textvariable=self.output_var, width=42).grid(
            row=0, column=0, padx=5, pady=4)
        tk.Button(out_frame, text="Save As…",
                  command=self._browse_output).grid(row=0, column=1, padx=5, pady=4)
        self.open_after_var = tk.BooleanVar(value=True)
        tk.Checkbutton(out_frame, text="Open in Word when done",
                       variable=self.open_after_var).grid(
            row=1, column=0, columnspan=2, sticky="w", padx=5, pady=(0, 4))

        # Per-page
        mode = tk.LabelFrame(self, text="Photos per Page")
        mode.grid(row=4, column=0, columnspan=2, sticky="ew", **pad)
        self.per_page_var = tk.IntVar(value=6)
        tk.Radiobutton(mode, text="4 per page (2×2)",
                       variable=self.per_page_var, value=4).grid(
            row=0, column=0, padx=10, pady=4, sticky="w")
        tk.Radiobutton(mode, text="6 per page (2×3)",
                       variable=self.per_page_var, value=6).grid(
            row=0, column=1, padx=10, pady=4, sticky="w")

        # Progress + status
        self.progress = ttk.Progressbar(self, length=440, mode="determinate")
        self.progress.grid(row=5, column=0, columnspan=2, padx=10, pady=(8, 2))
        self.status_var = tk.StringVar(value="Ready.")
        tk.Label(self, textvariable=self.status_var, anchor="w").grid(
            row=6, column=0, columnspan=2, sticky="ew", padx=10)

        # Generate
        self.gen_btn = tk.Button(self, text="Generate Contact Sheet",
                                 command=self._generate, width=30)
        self.gen_btn.grid(row=7, column=0, columnspan=2, pady=(4, 12))

    def _build_image_panel(self):
        dnd_hint = "  —  drag a folder or files here" if HAS_DND else ""
        frame = tk.LabelFrame(self, text=f"Images{dnd_hint}")

        # Buttons row
        btn_row = tk.Frame(frame)
        btn_row.pack(fill="x", padx=5, pady=(5, 2))
        tk.Button(btn_row, text="Browse Folder…",
                  command=self._browse_folder).pack(side="left", padx=(0, 4))
        tk.Button(btn_row, text="Add Files…",
                  command=self._browse_files).pack(side="left", padx=(0, 4))
        tk.Button(btn_row, text="Clear",
                  command=self._clear_images).pack(side="left", padx=(0, 12))
        self.subfolder_var = tk.BooleanVar(value=False)
        tk.Checkbutton(btn_row, text="Include sub-folders",
                       variable=self.subfolder_var).pack(side="left")

        # Native Listbox — fast even with 200+ files, EXTENDED selection = click to toggle
        list_outer = tk.Frame(frame, relief="sunken", bd=1)
        list_outer.pack(fill="both", expand=True, padx=5, pady=(2, 0))
        self._listbox = tk.Listbox(list_outer, selectmode=tk.EXTENDED,
                                   height=10, activestyle="none",
                                   exportselection=False)
        vsb = tk.Scrollbar(list_outer, orient="vertical",
                           command=self._listbox.yview)
        self._listbox.configure(yscrollcommand=vsb.set)
        self._listbox.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        self._listbox.bind("<<ListboxSelect>>", lambda _: self._update_count())

        # Count row
        count_row = tk.Frame(frame)
        count_row.pack(fill="x", padx=5, pady=(2, 5))
        self.count_var = tk.StringVar(value="No images added.")
        tk.Label(count_row, textvariable=self.count_var, anchor="w").pack(
            side="left", expand=True, fill="x")
        tk.Button(count_row, text="None", command=self._select_none,
                  width=5).pack(side="right", padx=(2, 0))
        tk.Button(count_row, text="All", command=self._select_all,
                  width=5).pack(side="right")

        return frame

    # ── Image list management ─────────────────────────────────────────────────

    def _add_images(self, paths):
        start = len(self._image_order)
        for path in paths:
            if path in self._image_set:
                continue
            self._image_order.append(path)
            self._image_set.add(path)
            self._listbox.insert(tk.END, os.path.basename(path))
        # Auto-select newly added items
        if len(self._image_order) > start:
            self._listbox.select_set(start, tk.END)
        self._update_count()

    def _collect_from_folder(self, folder):
        exts = {'.jpg', '.jpeg', '.png'}
        images = []
        if self.subfolder_var.get():
            for root, dirs, files in os.walk(folder):
                dirs.sort(key=natural_sort_key)
                for fname in sorted(files, key=natural_sort_key):
                    if os.path.splitext(fname)[1].lower() in exts:
                        images.append(os.path.join(root, fname))
        else:
            for fname in sorted(os.listdir(folder), key=natural_sort_key):
                if os.path.splitext(fname)[1].lower() in exts:
                    images.append(os.path.join(folder, fname))
        return images

    def _clear_images(self):
        self._listbox.delete(0, tk.END)
        self._image_order.clear()
        self._image_set.clear()
        self._update_count()

    def _select_all(self):
        self._listbox.select_set(0, tk.END)
        self._update_count()

    def _select_none(self):
        self._listbox.select_clear(0, tk.END)
        self._update_count()

    def _update_count(self):
        total = len(self._image_order)
        sel = len(self._listbox.curselection())
        if total == 0:
            self.count_var.set("No images added.")
        elif sel == total:
            self.count_var.set(f"{total} image{'s' if total != 1 else ''} selected.")
        else:
            self.count_var.set(f"{sel} of {total} images selected.")

    def _get_selected(self):
        return [self._image_order[i] for i in self._listbox.curselection()]

    # ── Event handlers ────────────────────────────────────────────────────────

    def _handle_drop(self, event):
        paths = self._parse_dnd(event.data)
        images = []
        for path in paths:
            if os.path.isdir(path):
                imgs = self._collect_from_folder(path)
                images.extend(imgs)
                if imgs:
                    self._auto_name_output(path)
            elif os.path.isfile(path):
                if os.path.splitext(path)[1].lower() in {'.jpg', '.jpeg', '.png'}:
                    images.append(path)
        if images:
            self._add_images(images)

    @staticmethod
    def _parse_dnd(raw):
        paths, i = [], 0
        while i < len(raw):
            if raw[i] == '{':
                end = raw.index('}', i)
                paths.append(raw[i + 1:end])
                i = end + 2
            elif raw[i] == ' ':
                i += 1
            else:
                end = raw.find(' ', i)
                if end == -1:
                    end = len(raw)
                paths.append(raw[i:end])
                i = end + 1
        return [p for p in paths if p]

    def _browse_folder(self):
        self.update_idletasks()
        folder = filedialog.askdirectory(title="Select Image Folder")
        if not folder:
            return
        images = self._collect_from_folder(folder)
        if images:
            self._add_images(images)
            self._auto_name_output(folder)
        else:
            messagebox.showwarning("No Images",
                                   "No JPG or PNG images found in that folder.")

    def _browse_files(self):
        self.update_idletasks()
        paths = filedialog.askopenfilenames(
            title="Select Images",
            filetypes=[("Image files", "*.jpg *.jpeg *.png"), ("All files", "*.*")])
        if paths:
            self._add_images(list(paths))

    def _auto_name_output(self, folder):
        name = os.path.basename(folder.rstrip('/\\'))
        parent = os.path.dirname(folder.rstrip('/\\'))
        self.output_var.set(os.path.join(parent, f"{name}_contact_sheet.docx"))

    def _browse_output(self):
        path = filedialog.asksaveasfilename(
            title="Save Contact Sheet As",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")])
        if path:
            self.output_var.set(path)

    def _generate(self):
        images = self._get_selected()
        output = self.output_var.get().strip()
        per_page = self.per_page_var.get()
        title = self.title_var.get().strip()

        if not images:
            messagebox.showerror("Error",
                                 "No images selected. Add images and check at least one.")
            return
        if not output:
            messagebox.showerror("Error", "Please specify an output file path.")
            return

        self.gen_btn.config(state="disabled")
        self.progress["value"] = 0
        self.status_var.set("Starting…")

        def run():
            try:
                build_contact_sheet(
                    images, output, per_page, title,
                    progress_cb=lambda v: self.after(
                        0, lambda: self.progress.config(value=v)),
                    status_cb=lambda s: self.after(
                        0, lambda: self.status_var.set(s)))
                if self.open_after_var.get():
                    try:
                        os.startfile(output)
                    except Exception:
                        pass
                self.after(0, lambda: messagebox.showinfo(
                    "Done", f"Contact sheet saved to:\n{output}"))
            except PermissionError:
                self.after(0, lambda: messagebox.showerror(
                    "Permission Error",
                    f"Cannot write to:\n{output}\n\n"
                    "Check that the file is not open in another program."))
            except Exception as e:
                err = str(e)
                self.after(0, lambda: messagebox.showerror(
                    "Error", f"Failed to generate contact sheet:\n{err}"))
            finally:
                self.after(0, lambda: self.gen_btn.config(state="normal"))

        threading.Thread(target=run, daemon=True).start()


if __name__ == "__main__":
    app = App()
    app.mainloop()
