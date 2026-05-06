"""
Word Search GUI — searches for keywords across .docx files in a selected folder.
Requires: python-docx
Build to .exe: see README_BUILD.md
"""

import csv
import os
import queue
import re
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    messagebox.showerror(
        "Missing dependency",
        "python-docx is not installed.\nRun: pip install python-docx",
    )
    raise SystemExit(1)


# ---------------------------------------------------------------------------
# Column layout (character widths for Labels; used by header + rows)
# ---------------------------------------------------------------------------
_COL_OPEN   = 26   # "Open in Word" button  (filename)
_COL_PATH   = 36   # full path
_COL_KW     = 18   # matched keyword
_COL_EXCPT  = 48   # excerpt


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _sentences_around(text: str, keyword: str, case_sensitive: bool, context: int = 120) -> str:
    flags = 0 if case_sensitive else re.IGNORECASE
    match = re.search(re.escape(keyword), text, flags=flags)
    if not match:
        return ""
    start = max(0, match.start() - context)
    end   = min(len(text), match.end() + context)
    excerpt = text[start:end].replace("\n", " ").strip()
    if start > 0:
        excerpt = "…" + excerpt
    if end < len(text):
        excerpt = excerpt + "…"
    return excerpt


def _extract_text(path: str) -> str:
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _collect_docx(folder: str, recursive: bool):
    if recursive:
        for root, _dirs, files in os.walk(folder):
            for f in files:
                if f.lower().endswith(".docx"):
                    yield os.path.join(root, f)
    else:
        for f in os.listdir(folder):
            if f.lower().endswith(".docx"):
                yield os.path.join(folder, f)


# ---------------------------------------------------------------------------
# Scrollable frame widget
# ---------------------------------------------------------------------------

class _ScrollableFrame(tk.Frame):
    """Vertically (and horizontally) scrollable container for arbitrary widgets."""

    def __init__(self, master, **kw):
        super().__init__(master, **kw)

        self._canvas = tk.Canvas(self, bg="#ffffff", highlightthickness=0)
        vsb = ttk.Scrollbar(self, orient="vertical",   command=self._canvas.yview)
        hsb = ttk.Scrollbar(self, orient="horizontal", command=self._canvas.xview)
        self._canvas.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self._canvas.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        self.rowconfigure(0, weight=1)
        self.columnconfigure(0, weight=1)

        self.inner = tk.Frame(self._canvas, bg="#ffffff")
        self._win_id = self._canvas.create_window((0, 0), window=self.inner, anchor="nw")

        self.inner.bind("<Configure>", self._on_inner_configure)
        self._canvas.bind("<Configure>", self._on_canvas_configure)
        # Mouse-wheel scrolling — only active when pointer is over this widget
        self._canvas.bind("<Enter>", self._bind_mousewheel)
        self._canvas.bind("<Leave>", self._unbind_mousewheel)

    def _on_inner_configure(self, _e):
        self._canvas.configure(scrollregion=self._canvas.bbox("all"))

    def _on_canvas_configure(self, event):
        self._canvas.itemconfig(self._win_id, width=event.width)

    def _bind_mousewheel(self, _e):
        self._canvas.bind_all("<MouseWheel>", self._on_mousewheel)

    def _unbind_mousewheel(self, _e):
        self._canvas.unbind_all("<MouseWheel>")

    def _on_mousewheel(self, event):
        self._canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def clear(self):
        for child in self.inner.winfo_children():
            child.destroy()


# ---------------------------------------------------------------------------
# Search worker (background thread)
# ---------------------------------------------------------------------------

def search_worker(
    folder: str,
    recursive: bool,
    keywords: list,
    match_all: bool,
    case_sensitive: bool,
    result_queue: queue.Queue,
):
    paths = list(_collect_docx(folder, recursive))
    total = len(paths)
    match_count = 0
    files_with_matches = set()

    for idx, path in enumerate(paths, 1):
        filename = os.path.basename(path)
        result_queue.put({"type": "progress", "current": idx, "total": total, "filename": filename})

        try:
            text = _extract_text(path)
        except Exception as exc:
            result_queue.put({"type": "error", "filename": filename, "path": path, "reason": str(exc)})
            continue

        flags = 0 if case_sensitive else re.IGNORECASE
        matched_keywords = [kw for kw in keywords if re.search(re.escape(kw), text, flags)]
        hit = (len(matched_keywords) == len(keywords)) if match_all else bool(matched_keywords)

        if hit:
            for kw in matched_keywords:
                excerpt = _sentences_around(text, kw, case_sensitive)
                result_queue.put({
                    "type": "match",
                    "filename": filename,
                    "path": path,
                    "keyword": kw,
                    "excerpt": excerpt,
                })
                match_count += 1
            files_with_matches.add(path)

    result_queue.put({"type": "done", "matches": match_count, "files_with_matches": files_with_matches})


# ---------------------------------------------------------------------------
# Main application
# ---------------------------------------------------------------------------

class WordSearchApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Word Search — .docx Files")
        self.minsize(960, 620)
        self.configure(bg="#f0f0f0")

        self._result_queue  = queue.Queue()
        self._search_thread = None
        self._all_rows      = []   # (filename, path, keyword, excerpt)
        self._error_rows    = []   # (filename, path, reason)
        self._row_count     = 0    # for alternating row colours

        self._build_ui()
        self.after(100, self._poll_queue)

    # ------------------------------------------------------------------
    # UI construction
    # ------------------------------------------------------------------

    def _build_ui(self):
        pad = {"padx": 8, "pady": 4}

        # ---- Settings panel ----
        top = tk.LabelFrame(self, text="Search Settings", bg="#f0f0f0",
                            font=("Segoe UI", 9, "bold"))
        top.pack(fill="x", padx=10, pady=(10, 4))

        tk.Label(top, text="Folder:", bg="#f0f0f0").grid(row=0, column=0, sticky="e", **pad)
        self._folder_var = tk.StringVar()
        tk.Entry(top, textvariable=self._folder_var, width=60).grid(row=0, column=1, sticky="ew", **pad)
        tk.Button(top, text="Browse…", command=self._browse_folder).grid(row=0, column=2, **pad)

        tk.Label(top, text="Keywords:", bg="#f0f0f0").grid(row=1, column=0, sticky="e", **pad)
        self._keywords_var = tk.StringVar()
        tk.Entry(top, textvariable=self._keywords_var, width=60).grid(row=1, column=1, sticky="ew", **pad)
        tk.Label(top, text="(comma-separated)", bg="#f0f0f0", fg="#666").grid(
            row=1, column=2, sticky="w", **pad)

        opts = tk.Frame(top, bg="#f0f0f0")
        opts.grid(row=2, column=0, columnspan=3, sticky="w", padx=8, pady=4)

        self._recursive_var = tk.BooleanVar(value=True)
        tk.Checkbutton(opts, text="Include subfolders", variable=self._recursive_var,
                       bg="#f0f0f0").pack(side="left", padx=(0, 16))

        self._match_all_var = tk.BooleanVar(value=False)
        tk.Checkbutton(opts, text="Match ALL keywords (AND)", variable=self._match_all_var,
                       bg="#f0f0f0").pack(side="left", padx=(0, 16))

        self._case_var = tk.BooleanVar(value=False)
        tk.Checkbutton(opts, text="Case-sensitive", variable=self._case_var,
                       bg="#f0f0f0").pack(side="left")

        top.columnconfigure(1, weight=1)

        # ---- Action bar ----
        action = tk.Frame(self, bg="#f0f0f0")
        action.pack(fill="x", padx=10, pady=2)

        self._search_btn = tk.Button(
            action, text="Search", bg="#0078d4", fg="white",
            font=("Segoe UI", 10, "bold"), padx=16, pady=4,
            relief="flat", cursor="hand2", command=self._start_search,
        )
        self._search_btn.pack(side="left")

        self._export_btn = tk.Button(
            action, text="Export to CSV…", padx=12, pady=4,
            relief="flat", cursor="hand2", command=self._export_csv, state="disabled",
        )
        self._export_btn.pack(side="left", padx=(8, 0))

        self._errors_btn = tk.Button(
            action, text="Show Skipped Files", padx=12, pady=4,
            relief="flat", cursor="hand2", command=self._show_errors, state="disabled",
        )
        self._errors_btn.pack(side="left", padx=(8, 0))

        self._status_var = tk.StringVar(value="Ready.")
        tk.Label(action, textvariable=self._status_var, bg="#f0f0f0", fg="#444",
                 font=("Segoe UI", 9)).pack(side="left", padx=16)

        # ---- Progress bar ----
        self._progress_var = tk.DoubleVar(value=0)
        ttk.Progressbar(self, variable=self._progress_var, maximum=100).pack(
            fill="x", padx=10, pady=(2, 4))

        # ---- Results section ----
        results_frame = tk.LabelFrame(self, text="Results", bg="#f0f0f0",
                                      font=("Segoe UI", 9, "bold"))
        results_frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        results_frame.rowconfigure(1, weight=1)
        results_frame.columnconfigure(0, weight=1)

        # Column header row
        header = tk.Frame(results_frame, bg="#c8d4e8")
        header.grid(row=0, column=0, sticky="ew")

        _hfont = ("Segoe UI", 9, "bold")
        tk.Label(header, text="Open in Word", font=_hfont, bg="#c8d4e8",
                 width=_COL_OPEN, anchor="w").pack(side="left", padx=(6, 2), pady=3)
        tk.Label(header, text="Full Path", font=_hfont, bg="#c8d4e8",
                 width=_COL_PATH, anchor="w").pack(side="left", padx=2, pady=3)
        tk.Label(header, text="Matched Keyword", font=_hfont, bg="#c8d4e8",
                 width=_COL_KW, anchor="w").pack(side="left", padx=2, pady=3)
        tk.Label(header, text="Excerpt", font=_hfont, bg="#c8d4e8",
                 width=_COL_EXCPT, anchor="w").pack(side="left", padx=2, pady=3)

        # Scrollable body
        self._results_sf = _ScrollableFrame(results_frame, bg="#ffffff")
        self._results_sf.grid(row=1, column=0, sticky="nsew")

    # ------------------------------------------------------------------
    # Row builder
    # ------------------------------------------------------------------

    def _add_result_row(self, filename: str, path: str, keyword: str, excerpt: str):
        self._row_count += 1
        bg = "#ffffff" if self._row_count % 2 == 1 else "#f0f4fa"

        row = tk.Frame(self._results_sf.inner, bg=bg)
        row.pack(fill="x")

        # --- Bold, highlighted "Open in Word" button with the filename ---
        btn = tk.Button(
            row,
            text=filename,
            font=("Segoe UI", 9, "bold"),
            bg="#FFE066",          # vivid yellow highlight
            fg="#1a237e",          # dark-navy text
            activebackground="#FFD600",
            activeforeground="#1a237e",
            relief="flat",
            bd=0,
            cursor="hand2",
            anchor="w",
            width=_COL_OPEN,
            padx=4,
            command=lambda p=path: self._open_in_word(p),
        )
        btn.pack(side="left", padx=(4, 2), pady=2)

        _lfont = ("Segoe UI", 9)
        tk.Label(row, text=path,    font=_lfont, bg=bg, width=_COL_PATH,  anchor="w",
                 wraplength=260).pack(side="left", padx=2, pady=2)
        tk.Label(row, text=keyword, font=_lfont, bg=bg, width=_COL_KW,   anchor="w").pack(
                 side="left", padx=2, pady=2)
        tk.Label(row, text=excerpt, font=_lfont, bg=bg, width=_COL_EXCPT, anchor="w",
                 wraplength=340, justify="left").pack(side="left", padx=2, pady=2)

    def _open_in_word(self, path: str):
        if os.path.exists(path):
            os.startfile(path)   # Windows: opens with the default .docx handler (Word)
        else:
            messagebox.showwarning("File not found", f"Cannot open:\n{path}")

    # ------------------------------------------------------------------
    # Search control
    # ------------------------------------------------------------------

    def _browse_folder(self):
        folder = filedialog.askdirectory(title="Select folder to search")
        if folder:
            self._folder_var.set(folder)

    def _start_search(self):
        folder = self._folder_var.get().strip()
        if not folder or not os.path.isdir(folder):
            messagebox.showwarning("No folder", "Please select a valid folder first.")
            return

        raw_kw = self._keywords_var.get().strip()
        if not raw_kw:
            messagebox.showwarning("No keywords", "Please enter at least one keyword.")
            return

        keywords = [k.strip() for k in raw_kw.split(",") if k.strip()]
        if not keywords:
            messagebox.showwarning("No keywords", "Could not parse any keywords.")
            return

        if self._search_thread and self._search_thread.is_alive():
            messagebox.showinfo("Busy", "A search is already running. Please wait.")
            return

        # Reset
        self._all_rows.clear()
        self._error_rows.clear()
        self._row_count = 0
        self._results_sf.clear()
        self._progress_var.set(0)
        self._export_btn.config(state="disabled")
        self._errors_btn.config(state="disabled")
        self._search_btn.config(state="disabled")
        self._status_var.set("Scanning…")

        self._search_thread = threading.Thread(
            target=search_worker,
            args=(folder, self._recursive_var.get(), keywords,
                  self._match_all_var.get(), self._case_var.get(), self._result_queue),
            daemon=True,
        )
        self._search_thread.start()

    # ------------------------------------------------------------------
    # Queue polling (runs on the main thread every 50 ms)
    # ------------------------------------------------------------------

    def _poll_queue(self):
        try:
            while True:
                msg   = self._result_queue.get_nowait()
                mtype = msg["type"]

                if mtype == "progress":
                    pct = (msg["current"] / msg["total"] * 100) if msg["total"] else 0
                    self._progress_var.set(pct)
                    self._status_var.set(
                        f"Scanning {msg['current']}/{msg['total']}: {msg['filename']}"
                    )

                elif mtype == "match":
                    self._all_rows.append(
                        (msg["filename"], msg["path"], msg["keyword"], msg["excerpt"])
                    )
                    self._add_result_row(
                        msg["filename"], msg["path"], msg["keyword"], msg["excerpt"]
                    )

                elif mtype == "error":
                    self._error_rows.append((msg["filename"], msg["path"], msg["reason"]))

                elif mtype == "done":
                    n_matches = msg["matches"]
                    n_files   = len(msg["files_with_matches"])
                    self._status_var.set(
                        f"Done. Found {n_matches} match{'es' if n_matches != 1 else ''}"
                        f" across {n_files} file{'s' if n_files != 1 else ''}."
                    )
                    self._progress_var.set(100)
                    self._search_btn.config(state="normal")
                    if self._all_rows:
                        self._export_btn.config(state="normal")
                    if self._error_rows:
                        self._errors_btn.config(state="normal")

        except queue.Empty:
            pass

        self.after(50, self._poll_queue)

    # ------------------------------------------------------------------
    # Export / error popup
    # ------------------------------------------------------------------

    def _export_csv(self):
        if not self._all_rows:
            messagebox.showinfo("Nothing to export", "No results to export.")
            return
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")],
            title="Save results as CSV",
        )
        if not dest:
            return
        try:
            with open(dest, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                writer.writerow(["File Name", "Full Path", "Matched Keyword", "Excerpt"])
                writer.writerows(self._all_rows)
            messagebox.showinfo("Exported", f"Results saved to:\n{dest}")
        except OSError as exc:
            messagebox.showerror("Export failed", str(exc))

    def _show_errors(self):
        if not self._error_rows:
            return
        win = tk.Toplevel(self)
        win.title("Skipped / Unreadable Files")
        win.minsize(700, 300)

        cols = ("filename", "path", "reason")
        tree = ttk.Treeview(win, columns=cols, show="headings")
        tree.heading("filename", text="File Name")
        tree.heading("path",     text="Full Path")
        tree.heading("reason",   text="Reason Skipped")
        tree.column("filename", width=160)
        tree.column("path",     width=280)
        tree.column("reason",   width=240)

        vsb = ttk.Scrollbar(win, orient="vertical", command=tree.yview)
        tree.configure(yscrollcommand=vsb.set)
        tree.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")

        for row in self._error_rows:
            tree.insert("", "end", values=row)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    app = WordSearchApp()
    app.mainloop()
