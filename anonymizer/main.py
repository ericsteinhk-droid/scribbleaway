#!/usr/bin/env python3
"""
Office File Anonymizer
  - Person names  ->  initials        (e.g. "John Smith" -> "J.S.")
  - Org names     ->  numbered labels (e.g. "Acme Corp"  -> "org1")
Supports .docx (Word) and .xlsx (Excel).
Output is saved alongside the original with the suffix _anon.
"""

import sys
import os
import threading
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, ttk, scrolledtext, messagebox

# ---------------------------------------------------------------------------
# NLP setup – tries the direct package import first (PyInstaller-safe),
# then falls back to spacy.load() for dev environments.
# ---------------------------------------------------------------------------
_nlp = None


def _load_nlp():
    global _nlp
    try:
        import en_core_web_sm
        _nlp = en_core_web_sm.load()
    except Exception:
        try:
            import spacy
            _nlp = spacy.load("en_core_web_sm")
        except Exception:
            pass


_load_nlp()


# ---------------------------------------------------------------------------
# Anonymizer – tracks entity->replacement mappings across all processed files
# ---------------------------------------------------------------------------

class Anonymizer:
    def __init__(self):
        self.person_map: dict = {}   # "John Smith" -> "J.S."
        self.org_map: dict = {}      # "Acme Corp"  -> "org1"
        self._org_counter = 0

    @staticmethod
    def _initials(name: str) -> str:
        return "".join(w[0].upper() + "." for w in name.split() if w)

    def _org_label(self, name: str) -> str:
        key = name.strip()
        if key not in self.org_map:
            self._org_counter += 1
            self.org_map[key] = f"org{self._org_counter}"
        return self.org_map[key]

    def get_replacements(self, text: str) -> list:
        """
        Run NER on *text* and return [(original, replacement), ...] sorted
        longest-first. Side-effect: updates person_map / org_map.
        """
        if not _nlp or not text or not text.strip():
            return []
        doc = _nlp(text)
        seen: set = set()
        pairs = []
        for ent in doc.ents:
            if ent.text in seen:
                continue
            seen.add(ent.text)
            if ent.label_ == "PERSON":
                init = self._initials(ent.text)
                self.person_map.setdefault(ent.text, init)
                pairs.append((ent.text, init))
            elif ent.label_ == "ORG":
                pairs.append((ent.text, self._org_label(ent.text)))
        pairs.sort(key=lambda x: len(x[0]), reverse=True)
        return pairs

    def anonymize(self, text: str) -> str:
        """Anonymize a plain string using character-offset substitution."""
        if not _nlp or not text or not text.strip():
            return text
        doc = _nlp(text)
        changes = []
        seen: set = set()
        for ent in doc.ents:
            span = (ent.start_char, ent.end_char)
            if span in seen:
                continue
            seen.add(span)
            if ent.label_ == "PERSON":
                init = self._initials(ent.text)
                self.person_map.setdefault(ent.text, init)
                changes.append((ent.start_char, ent.end_char, init))
            elif ent.label_ == "ORG":
                changes.append((ent.start_char, ent.end_char,
                                 self._org_label(ent.text)))
        # Apply right-to-left so earlier char positions stay valid
        for start, end, repl in sorted(changes, reverse=True):
            text = text[:start] + repl + text[end:]
        return text


# ---------------------------------------------------------------------------
# Word (.docx) processing
# ---------------------------------------------------------------------------

def _anonymize_paragraph(para, anon: Anonymizer) -> None:
    runs = list(para.runs)
    if not runs:
        return
    full = "".join(r.text for r in runs)
    if not full.strip():
        return

    replacements = anon.get_replacements(full)
    if not replacements:
        return

    active = list(runs)
    for original, replacement in replacements:
        # Attempt single-run replacement first to preserve inline formatting
        replaced = False
        for run in active:
            if original in run.text:
                run.text = run.text.replace(original, replacement)
                replaced = True

        if not replaced:
            # Entity spans multiple runs: merge into first, blank the rest
            combined = "".join(r.text for r in active)
            if original in combined:
                active[0].text = combined.replace(original, replacement)
                for r in active[1:]:
                    r.text = ""
                active = [active[0]]


def process_docx(path: str, anon: Anonymizer, log) -> str:
    from docx import Document

    doc = Document(path)

    for para in doc.paragraphs:
        _anonymize_paragraph(para, anon)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _anonymize_paragraph(para, anon)

    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer):
            try:
                for para in hf.paragraphs:
                    _anonymize_paragraph(para, anon)
            except Exception:
                pass

    out = _anon_path(path)
    doc.save(out)
    log(f"  Saved: {out}")
    return out


# ---------------------------------------------------------------------------
# Excel (.xlsx) processing
# ---------------------------------------------------------------------------

def process_xlsx(path: str, anon: Anonymizer, log) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(path)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.strip():
                    new_val = anon.anonymize(cell.value)
                    if new_val != cell.value:
                        cell.value = new_val

    out = _anon_path(path)
    wb.save(out)
    log(f"  Saved: {out}")
    return out


def _anon_path(path: str) -> str:
    p = Path(path)
    return str(p.parent / (p.stem + "_anon" + p.suffix))


# ---------------------------------------------------------------------------
# GUI
# ---------------------------------------------------------------------------

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Office File Anonymizer")
        self.geometry("720x580")
        self.minsize(520, 440)
        ttk.Style(self).theme_use("clam")
        self._build_ui()
        if not _nlp:
            messagebox.showerror(
                "NLP model not found",
                "The spaCy English model could not be loaded.\n\n"
                "Please run:\n    python -m spacy download en_core_web_sm\n\n"
                "Then restart the application.",
                parent=self,
            )

    # ── UI construction ──────────────────────────────────────────────────────

    def _build_ui(self):
        PAD = dict(padx=10, pady=4)

        # Header
        hdr = ttk.Frame(self, padding=(10, 10, 10, 2))
        hdr.pack(fill=tk.X)
        ttk.Label(hdr, text="Office File Anonymizer",
                  font=("Segoe UI", 13, "bold")).pack(anchor=tk.W)
        ttk.Label(hdr,
                  text="Replaces names with initials and org names with org1, org2, ...",
                  font=("Segoe UI", 9), foreground="#555").pack(anchor=tk.W)
        ttk.Separator(self).pack(fill=tk.X, padx=10, pady=4)

        # File list label
        ttk.Label(self, text="Files to process:", font=("Segoe UI", 10, "bold"),
                  padding=(10, 0)).pack(anchor=tk.W)

        # File listbox
        lf = ttk.Frame(self, padding=(10, 2, 10, 0))
        lf.pack(fill=tk.BOTH, expand=True)
        sb = ttk.Scrollbar(lf)
        sb.pack(side=tk.RIGHT, fill=tk.Y)
        self.lb = tk.Listbox(lf, yscrollcommand=sb.set,
                              selectmode=tk.EXTENDED,
                              font=("Consolas", 9), height=10,
                              activestyle="dotbox")
        self.lb.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        sb.config(command=self.lb.yview)

        # File buttons
        bf = ttk.Frame(self, padding=(10, 4))
        bf.pack(fill=tk.X)
        ttk.Button(bf, text="Add Files...",  command=self._add_files).pack(side=tk.LEFT)
        ttk.Button(bf, text="Add Folder...", command=self._add_folder).pack(side=tk.LEFT, padx=6)
        ttk.Button(bf, text="Remove",        command=self._remove).pack(side=tk.LEFT)
        ttk.Button(bf, text="Clear All",     command=self._clear).pack(side=tk.LEFT, padx=6)

        # Progress
        pf = ttk.Frame(self, padding=(10, 2))
        pf.pack(fill=tk.X)
        self.status_lbl = ttk.Label(pf, text="Ready", foreground="#444")
        self.status_lbl.pack(anchor=tk.W)
        self.progress = ttk.Progressbar(pf, mode="determinate")
        self.progress.pack(fill=tk.X, pady=2)

        # Log
        log_frame = ttk.LabelFrame(self, text="Log", padding=(6, 4))
        log_frame.pack(fill=tk.BOTH, expand=True, **PAD)
        self.log_txt = scrolledtext.ScrolledText(
            log_frame, state=tk.DISABLED, font=("Consolas", 9), height=8)
        self.log_txt.pack(fill=tk.BOTH, expand=True)

        # Run button
        self.run_btn = ttk.Button(self, text="  Anonymize Files  ",
                                   command=self._start)
        self.run_btn.pack(pady=(2, 12))

    # ── File list management ─────────────────────────────────────────────────

    def _add_files(self):
        chosen = filedialog.askopenfilenames(
            title="Select Word or Excel files",
            filetypes=[("Office files", "*.docx *.xlsx"),
                       ("Word files",   "*.docx"),
                       ("Excel files",  "*.xlsx")])
        existing = set(self.lb.get(0, tk.END))
        for f in chosen:
            if f not in existing:
                self.lb.insert(tk.END, f)
                existing.add(f)

    def _add_folder(self):
        folder = filedialog.askdirectory(title="Select a folder")
        if not folder:
            return
        existing = set(self.lb.get(0, tk.END))
        for p in sorted(Path(folder).rglob("*")):
            if p.suffix.lower() in (".docx", ".xlsx") and "_anon" not in p.stem:
                s = str(p)
                if s not in existing:
                    self.lb.insert(tk.END, s)
                    existing.add(s)

    def _remove(self):
        for i in reversed(self.lb.curselection()):
            self.lb.delete(i)

    def _clear(self):
        self.lb.delete(0, tk.END)

    # ── Logging ──────────────────────────────────────────────────────────────

    def _log(self, msg: str):
        def _do():
            self.log_txt.config(state=tk.NORMAL)
            self.log_txt.insert(tk.END, msg + "\n")
            self.log_txt.see(tk.END)
            self.log_txt.config(state=tk.DISABLED)
        self.after(0, _do)

    # ── Processing ───────────────────────────────────────────────────────────

    def _start(self):
        files = list(self.lb.get(0, tk.END))
        if not files:
            messagebox.showwarning("No files", "Add files or a folder first.")
            return
        if not _nlp:
            messagebox.showerror("Model unavailable",
                                  "spaCy NLP model is not loaded. "
                                  "Run:  python -m spacy download en_core_web_sm")
            return
        self.run_btn.config(state=tk.DISABLED)
        threading.Thread(target=self._worker, args=(files,), daemon=True).start()

    def _worker(self, files: list):
        anon = Anonymizer()
        n = len(files)
        self.after(0, lambda: self.progress.config(maximum=n, value=0))
        errors = 0

        for i, fp in enumerate(files, 1):
            name = Path(fp).name
            self._log(f"\n[{i}/{n}] {name}")
            self.after(0, lambda lbl=f"{name}  ({i}/{n})":
                       self.status_lbl.config(text=lbl))
            try:
                ext = Path(fp).suffix.lower()
                if ext == ".docx":
                    process_docx(fp, anon, self._log)
                elif ext == ".xlsx":
                    process_xlsx(fp, anon, self._log)
                else:
                    self._log("  Skipped (unsupported format)")
            except Exception as exc:
                self._log(f"  ERROR: {exc}")
                errors += 1
            self.after(0, lambda v=i: self.progress.config(value=v))

        summary = (
            f"\nDone.  {n} file(s) processed, {errors} error(s).\n"
            f"  Persons anonymized : {len(anon.person_map)}\n"
            f"  Orgs anonymized    : {len(anon.org_map)}"
        )
        self._log(summary)
        self.after(0, lambda: [
            self.status_lbl.config(text="Done"),
            self.run_btn.config(state=tk.NORMAL),
        ])


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    App().mainloop()
