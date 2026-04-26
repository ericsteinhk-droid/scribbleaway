#!/usr/bin/env python3
"""
Offline anonymizer for Word (.docx) and Excel (.xlsx) files.

Replaces personally identifiable information (PII) with stable placeholder
tokens before uploading documents to cloud services.  All processing happens
locally — no data ever leaves the machine.

Usage:
    python anonymize.py report.docx
    python anonymize.py data.xlsx --dates --no-nlp
    python anonymize.py *.docx *.xlsx --save-mapping
    python anonymize.py report.docx --output-dir ./clean/

Supported PII types (always on):
    EMAIL, PHONE, SSN, CREDIT_CARD, IP_ADDRESS, URL

Optional (requires --dates flag):
    DATE

Named-entity recognition (requires spaCy — see requirements.txt):
    PERSON, ORG, LOCATION
"""

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

_PATTERNS: Dict[str, re.Pattern] = {
    "EMAIL": re.compile(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b"),
    "PHONE": re.compile(
        r"(?<!\d)(?:\+?1[\s.\-]?)?(?:\(?\d{3}\)?[\s.\-]?)?\d{3}[\s.\-]\d{4}(?!\d)"
    ),
    "SSN": re.compile(r"\b\d{3}[\s\-]?\d{2}[\s\-]?\d{4}\b"),
    "CREDIT_CARD": re.compile(r"\b(?:\d[ \-]?){13,16}\b"),
    "IP_ADDRESS": re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b"),
    "URL": re.compile(r"https?://[^\s<>\"{}|\\^`\[\]]+"),
    "DATE": re.compile(
        r"\b(?:\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}"
        r"|\d{4}[/\-\.]\d{1,2}[/\-\.]\d{1,2}"
        r"|(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?"
        r"|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
        r"\.?\s+\d{1,2},?\s+\d{4})\b",
        re.IGNORECASE,
    ),
}

# spaCy entity labels → placeholder prefix
_NLP_LABELS = {"PERSON": "PERSON", "ORG": "ORG", "GPE": "LOCATION", "LOC": "LOCATION"}


# ---------------------------------------------------------------------------
# Core anonymizer
# ---------------------------------------------------------------------------

class Anonymizer:
    """Stateful anonymizer: repeated occurrences of the same string always
    receive the same placeholder token, enabling consistent de-anonymization."""

    def __init__(self, use_nlp: bool = True, anonymize_dates: bool = False):
        self.anonymize_dates = anonymize_dates
        self._mapping: Dict[str, str] = {}   # original  → placeholder
        self._counters: Dict[str, int] = {}
        self.nlp = None
        if use_nlp:
            self._init_nlp()

    # ── setup ──────────────────────────────────────────────────────────────

    def _init_nlp(self):
        try:
            import spacy  # noqa: F401
        except ImportError:
            print(
                "[warn] spaCy not installed — skipping named-entity recognition.\n"
                "       pip install spacy && python -m spacy download en_core_web_sm"
            )
            return
        try:
            import spacy
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            print(
                "[warn] spaCy model 'en_core_web_sm' not found — skipping NER.\n"
                "       python -m spacy download en_core_web_sm"
            )

    # ── placeholder management ─────────────────────────────────────────────

    def _token(self, kind: str, original: str) -> str:
        if original in self._mapping:
            return self._mapping[original]
        n = self._counters.get(kind, 0) + 1
        self._counters[kind] = n
        placeholder = f"[{kind}_{n}]"
        self._mapping[original] = placeholder
        return placeholder

    # ── text anonymization ─────────────────────────────────────────────────

    def anonymize_text(self, text: str) -> str:
        if not text or not isinstance(text, str):
            return text

        spans: List[Tuple[int, int, str]] = []

        # Named-entity recognition
        if self.nlp:
            doc = self.nlp(text)
            for ent in doc.ents:
                if ent.label_ in _NLP_LABELS:
                    spans.append(
                        (ent.start_char, ent.end_char,
                         self._token(_NLP_LABELS[ent.label_], ent.text))
                    )

        # Regex patterns
        active = {
            k: v for k, v in _PATTERNS.items()
            if k != "DATE" or self.anonymize_dates
        }
        for kind, pattern in active.items():
            for m in pattern.finditer(text):
                s, e = m.start(), m.end()
                # Skip if already covered by an NLP span
                if any(a <= s < b or a < e <= b for a, b, _ in spans):
                    continue
                spans.append((s, e, self._token(kind, m.group())))

        if not spans:
            return text

        # Apply replacements back-to-front so indices stay valid
        spans.sort(key=lambda x: x[0], reverse=True)
        result = text
        for s, e, token in spans:
            result = result[:s] + token + result[e:]
        return result

    # ── mapping persistence ────────────────────────────────────────────────

    def save_mapping(self, path: Path):
        data = {
            "note": "CONFIDENTIAL — maps placeholder tokens back to original values.",
            "mapping": {v: k for k, v in self._mapping.items()},
        }
        path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")

    # ── Word (.docx) ───────────────────────────────────────────────────────

    def anonymize_docx(self, src: Path, dst: Path) -> int:
        """Returns number of paragraphs modified."""
        from docx import Document  # type: ignore

        doc = Document(str(src))
        changed = 0

        def process(paragraphs):
            nonlocal changed
            for para in paragraphs:
                if self._anonymize_paragraph(para):
                    changed += 1

        process(doc.paragraphs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process(cell.paragraphs)

        for section in doc.sections:
            for hf in (
                section.header, section.footer,
                section.even_page_header, section.even_page_footer,
                section.first_page_header, section.first_page_footer,
            ):
                try:
                    process(hf.paragraphs)
                except Exception:
                    pass

        doc.save(str(dst))
        return changed

    def _anonymize_paragraph(self, para) -> bool:
        """Anonymize a paragraph in-place. Returns True if text changed."""
        original = para.text
        if not original.strip():
            return False
        anonymized = self.anonymize_text(original)
        if anonymized == original:
            return False
        # Collapse all runs into the first run so the replacement is correct.
        # The first run's character-level formatting (font, bold, etc.) is kept
        # for the whole paragraph — acceptable trade-off for anonymization.
        if para.runs:
            para.runs[0].text = anonymized
            for run in para.runs[1:]:
                run.text = ""
        return True

    # ── Excel (.xlsx) ──────────────────────────────────────────────────────

    def anonymize_xlsx(self, src: Path, dst: Path) -> int:
        """Returns number of cells modified."""
        import openpyxl  # type: ignore

        wb = openpyxl.load_workbook(str(src))
        changed = 0
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        new_val = self.anonymize_text(cell.value)
                        if new_val != cell.value:
                            cell.value = new_val
                            changed += 1
        wb.save(str(dst))
        return changed


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def build_output_path(src: Path, output_dir: Optional[Path]) -> Path:
    stem = src.stem + "_anonymized"
    name = stem + src.suffix
    base = output_dir if output_dir else src.parent
    return base / name


def process_file(path: Path, anon: Anonymizer, output_dir: Optional[Path]) -> Path:
    dst = build_output_path(path, output_dir)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        n = anon.anonymize_docx(path, dst)
        print(f"  {path.name} → {dst.name}  ({n} paragraph(s) changed)")
    elif suffix == ".xlsx":
        n = anon.anonymize_xlsx(path, dst)
        print(f"  {path.name} → {dst.name}  ({n} cell(s) changed)")
    else:
        raise ValueError(f"Unsupported file type: {suffix!r} (expected .docx or .xlsx)")

    return dst


def main():
    parser = argparse.ArgumentParser(
        description="Anonymize Word/Excel files offline before cloud upload.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("files", nargs="+", type=Path, metavar="FILE",
                        help=".docx or .xlsx file(s) to anonymize")
    parser.add_argument("--output-dir", type=Path, metavar="DIR",
                        help="Write anonymized files here (default: same directory as input)")
    parser.add_argument("--save-mapping", action="store_true",
                        help="Save a JSON file mapping placeholders back to originals")
    parser.add_argument("--dates", action="store_true",
                        help="Also anonymize dates (off by default)")
    parser.add_argument("--no-nlp", action="store_true",
                        help="Skip spaCy NER; use regex patterns only")
    args = parser.parse_args()

    if args.output_dir:
        args.output_dir.mkdir(parents=True, exist_ok=True)

    anon = Anonymizer(use_nlp=not args.no_nlp, anonymize_dates=args.dates)

    errors = []
    processed = []
    for path in args.files:
        if not path.exists():
            print(f"[skip] {path}: file not found")
            errors.append(path)
            continue
        try:
            dst = process_file(path, anon, args.output_dir)
            processed.append(dst)
        except Exception as exc:
            print(f"[error] {path}: {exc}")
            errors.append(path)

    if args.save_mapping and processed:
        # Save mapping next to the first output file
        map_path = processed[0].parent / "anonymization_mapping.json"
        anon.save_mapping(map_path)
        print(f"  Mapping saved → {map_path}")
        print("  Keep the mapping file confidential — it contains the original values.")

    total = len(args.files)
    ok = len(processed)
    print(f"\nDone: {ok}/{total} file(s) anonymized" + (f", {len(errors)} error(s)" if errors else "."))

    if errors:
        sys.exit(1)


if __name__ == "__main__":
    main()
